import os
import pdfplumber
from docx import Document
import aiofiles
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Get configuration from environment variables
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "storage/uploads")

async def save_upload_file(file):
    """Save an uploaded file to the uploads directory and return the path"""
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    file_path = f"{UPLOAD_DIR}/{file.filename}"
    async with aiofiles.open(file_path, 'wb') as out_file:
        content = await file.read()
        await out_file.write(content)
    return file_path

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file"""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
                text += "\n\n"  # Add page breaks
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        text = f"Error processing PDF: {str(e)}"
    return text

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        text = f"Error processing DOCX: {str(e)}"
    return text

def extract_text_from_file(file_path):
    """Extract text from a file based on its extension"""
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        return f"Unsupported file format: {file_ext}"

async def get_documents_list():
    """Get a list of all uploaded documents"""
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    files = []
    for file_name in os.listdir(UPLOAD_DIR):
        file_path = os.path.join(UPLOAD_DIR, file_name)
        if os.path.isfile(file_path):
            files.append({
                "name": file_name,
                "path": file_path,
                "size": os.path.getsize(file_path)
            })
    return files
