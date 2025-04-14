from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os
from pathlib import Path
import time
from datetime import datetime
from dotenv import load_dotenv
import glob

# For PDF conversion - temporarily disabled
# from weasyprint import HTML, CSS
# from weasyprint.text.fonts import FontConfiguration
import tempfile

# Load environment variables
load_dotenv()

# Get configuration from environment variables
OUTPUT_DIR = os.getenv("OUTPUT_DIR", "storage/outputs")

def create_docx(extracted_text, filename, document_type="emergency_plan"):
    """
    Create a well-formatted DOCX file from extracted text
    
    Args:
        extracted_text: The analyzed text content
        filename: Base filename for the output
        document_type: Type of document (for styling)
    
    Returns:
        Path to the saved DOCX file
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.author = "Document Analyzer"
    core_properties.title = filename.replace('_', ' ').title()
    core_properties.created = datetime.now()
    
    # Define custom styles
    styles = doc.styles
    
    # Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    heading1_style = styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Calibri'
    heading1_font.size = Pt(16)
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    heading1_style.paragraph_format.space_before = Pt(24)
    heading1_style.paragraph_format.space_after = Pt(18)
    
    # Heading 2 style
    heading2_style = styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Calibri'
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    heading2_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    heading2_style.paragraph_format.space_before = Pt(18)
    heading2_style.paragraph_format.space_after = Pt(12)
    
    # Title page
    title = doc.add_heading(document_type.replace('_', ' ').upper(), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(24)
    title.style.font.bold = True
    title.style.font.color.rgb = RGBColor(0, 51, 102)
    
    date_paragraph = doc.add_paragraph()
    date_paragraph.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph.style.font.italic = True
    
    # Ensure the table of contents is populated or remove it if not needed
    # If the table of contents is not being used, consider removing it
    # doc.add_page_break()  # After table of contents
    
    # Process the extracted text
    sections = extracted_text.split('\n# ')
    
    # Handle the first section (or intro text)
    if sections and not sections[0].startswith('# '):
        intro_text = sections[0]
        intro_para = doc.add_paragraph(intro_text)
        intro_para.style = normal_style
        sections.pop(0)
    
    # Process remaining sections
    for section in sections:
        if section.strip():
            lines = section.split('\n', 1)
            
            if len(lines) > 0:
                heading_text = lines[0].strip().replace('#', '').strip()
                heading = doc.add_heading(heading_text, 1)
                heading.style = heading1_style
                
                if len(lines) > 1:
                    content = lines[1]
                    
                    # Process subsections
                    subsections = content.split('\n## ')
                    
                    # Handle first subsection or content
                    if subsections:
                        first_part = subsections[0]
                        if not first_part.startswith('## '):
                            paragraphs = first_part.split('\n\n')
                            for para in paragraphs:
                                if para.strip():
                                    # Check if it's a bullet point list
                                    if '\n- ' in para or para.startswith('- '):
                                        bullet_items = para.split('\n- ')
                                        for i, item in enumerate(bullet_items):
                                            if i == 0 and not item.startswith('- '):
                                                p = doc.add_paragraph(item)
                                                p.style = normal_style
                                            else:
                                                bullet_text = item.replace('- ', '', 1) if item.startswith('- ') else item
                                                p = doc.add_paragraph(bullet_text, style='List Bullet')
                                                p.style.paragraph_format.left_indent = Inches(0.5)
                                                p.style.paragraph_format.space_after = Pt(6)
                                    else:
                                        p = doc.add_paragraph(para)
                                        p.style = normal_style
                    
                    # Process subsections
                    for i, subsection in enumerate(subsections):
                        if i == 0 and not subsection.startswith('## '):
                            continue
                        
                        sub_lines = subsection.split('\n', 1)
                        if len(sub_lines) > 0:
                            sub_heading = sub_lines[0].strip().replace('##', '').strip()
                            sub_heading = doc.add_heading(sub_heading, 2)
                            sub_heading.style = heading2_style
                            
                            if len(sub_lines) > 1:
                                sub_content = sub_lines[1]
                                sub_paras = sub_content.split('\n\n')
                                for para in sub_paras:
                                    if para.strip():
                                        if '\n- ' in para or para.startswith('- '):
                                            bullet_items = para.split('\n- ')
                                            for j, item in enumerate(bullet_items):
                                                if j == 0 and not item.startswith('- '):
                                                    p = doc.add_paragraph(item)
                                                    p.style = normal_style
                                                else:
                                                    bullet_text = item.replace('- ', '', 1) if item.startswith('- ') else item
                                                    p = doc.add_paragraph(bullet_text, style='List Bullet')
                                                    p.style.paragraph_format.left_indent = Inches(0.5)
                                                    p.style.paragraph_format.space_after = Pt(6)
                                        else:
                                            p = doc.add_paragraph(para)
                                            p.style = normal_style
    
    # Save document
    filename_safe = Path(filename).stem.replace(" ", "_")
    
    # Find the next available number in sequence
    existing_files = glob.glob(f"{OUTPUT_DIR}/{filename_safe}_*.docx")
    
    if not existing_files:
        next_number = 1
    else:
        # Extract numbers from existing filenames
        numbers = []
        for file in existing_files:
            base_name = os.path.basename(file)
            try:
                # Extract the number from filename_X.docx format
                number_str = base_name.replace(f"{filename_safe}_", "").replace(".docx", "")
                if number_str.isdigit():
                    numbers.append(int(number_str))
            except:
                continue
        
        next_number = max(numbers) + 1 if numbers else 1
    
    save_path = f"{OUTPUT_DIR}/{filename_safe}_{next_number}.docx"
    doc.save(save_path)
    return save_path

def create_pdf_from_docx(docx_path):
    """
    Convert a DOCX file to PDF
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Path to the generated PDF file (or None if PDF generation is disabled)
    """
    # PDF generation is temporarily disabled due to WeasyPrint dependencies
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    print("WARNING: PDF generation is temporarily disabled. Only DOCX output will be available.")
    return None

def generate_output_files(extracted_text, filename, formats=None):
    """
    Generate output files in the requested formats
    
    Args:
        extracted_text: Analyzed text content
        filename: Base filename
        formats: List of formats to generate (default: ["docx", "pdf"])
        
    Returns:
        Dictionary with paths to the generated files
    """
    if formats is None:
        formats = ["docx", "pdf"]
    
    output_files = {}
    
    if "docx" in formats:
        docx_path = create_docx(extracted_text, filename)
        output_files["docx"] = docx_path
    
    if "pdf" in formats and "docx" in output_files:
        # PDF generation is disabled, but we'll add the warning
        pdf_result = create_pdf_from_docx(output_files["docx"])
        if pdf_result:
            output_files["pdf"] = pdf_result
    
    return output_files
